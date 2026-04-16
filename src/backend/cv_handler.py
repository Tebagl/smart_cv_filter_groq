"""
================================================================================
SMART CV FILTER - DOCUMENT PROCESSING ENGINE (CV_HANDLER)
================================================================================
Author: Tebagl
Date: April 2026
Version: 2.0 (Cloud Migration)

Description:
The document processing layer responsible for extracting, cleaning, and 
standardizing text from various file formats (PDF, DOCX). This module ensures 
that the AI receives high-quality string data for accurate analysis.

TECHNICAL SPECIFICATIONS:
- Supported Formats: PDF (.pdf), Word (.docx)
- Libraries: PyMuPDF (fitz), PDFMiner, python-docx
- Cleaning Pipeline: Normalizes whitespace, removes non-printable characters, 
  and handles encoding issues.

KEY FUNCTIONS:
- extract_text(): Primary entry point for multi-format text recovery.
- clean_text(): Post-processing to optimize token usage for the LLM.
- get_file_metadata(): Optional tracking of file creation/modification dates.

ROBUSTNESS:
- Error handling for corrupted files and encrypted PDFs.
- Automatic detection of file extensions.
- Memory-efficient processing for large batches of documents.
================================================================================
"""

import os
import shutil
import logging
import re
import fitz  # PyMuPDF
import docx # .docx
from datetime import datetime
import csv

logger = logging.getLogger(__name__)

class CVHandler:
    def __init__(self, analyzer):
        """
        Manejador de CVs para sistema de carpetas e IA Embebida.
        """
        self.analyzer = analyzer
        # Inicializamos en None o vacío. 
        # No creamos carpetas aquí para evitar duplicados al arrancar.
        self.base_output = None 

    def _ensure_folders(self):
        """Crea la estructura de carpetas solo si base_output ha sido definido."""
        if not self.base_output:
            return

        os.makedirs(self.base_output, exist_ok=True)
        for folder in ["RECLUTADOS", "DESCARTADOS", "DUDAS"]:
            os.makedirs(os.path.join(self.base_output, folder), exist_ok=True)
    
    def _extract_text_from_pdf(self, pdf_path):
        """Extrae todo el texto de un archivo PDF."""
        text = ""
        try:
            with fitz.open(pdf_path) as doc:
                for page in doc:
                    text += page.get_text()
            return text
        except Exception as e:
            logger.error(f"Error leyendo PDF {pdf_path}: {e}")
            return ""
        
    def _extract_text_from_docx(self, docx_path):
        import sys
        print(f"\n>>> HILO VIVO: Intentando abrir {docx_path}", flush=True)
        
        try:
            from docx import Document
            print(">>> LIBRERÍA CARGADA CORRECTAMENTE", flush=True)
            doc = Document(docx_path)
            
            text_parts = []

            # 1. Extraer párrafos normales
            for para in doc.paragraphs:
                if para.text.strip():
                    text_parts.append(para.text)

            # 2. Extraer texto dentro de TABLAS (Crucial para CVs de internet)
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        # Limpiamos el texto de la celda
                        cell_text = cell.text.strip()
                        if cell_text and cell_text not in text_parts:
                            text_parts.append(cell_text)
            
            final_text = "\n".join(text_parts)
            
            # Verificación en terminal
            if final_text.strip():
                print(f"✅ ÉXITO: {len(final_text)} caracteres extraídos.", flush=True)
            else:
                print(f"⚠️ AVISO: El documento parece estar realmente vacío de texto.", flush=True)
                
            return final_text

        except Exception as e:
            print(f">>> ERROR DETECTADO: {e}", flush=True)
            return ""
        
    def _append_to_report(self, data):
        
        if not self.base_output: # <--- Seguridad 
            return
        
        """Añade una fila al archivo CSV de resumen."""
        report_path = os.path.join(self.base_output, "resumen_proceso.csv")
        file_exists = os.path.isfile(report_path)
        
        with open(report_path, mode='a', newline='', encoding='utf-8') as f:
            writer = csv.DictWriter(f, fieldnames=["Archivo", "Score", "Motivo", "Decision"])
            if not file_exists:
                writer.writeheader() # Si es nuevo, escribe la cabecera
            
            writer.writerow({
                "Archivo": os.path.basename(data["dest_path"]),
                "Score": data["score"],
                "Motivo": data["reason"],
                "Decision": data["decision"]
            })
        

    def process_cv(self, file_path: str, user_job_desc: str = None):
        """
        Analiza el CV y lo mueve físicamente según el score y guarda la razón.
        """
        try:
            ext = os.path.splitext(file_path)[1].lower()
            
            # --- Lógica Multiformato Ampliada ---
            if ext == ".pdf":
                raw_text = self._extract_text_from_pdf(file_path)
            elif ext == ".docx":
                raw_text = self._extract_text_from_docx(file_path)
            else:
                # Soporte para .txt y otros formatos de texto plano
                with open(file_path, 'r', encoding='utf-8') as f:
                    raw_text = f.read()

            if not raw_text.strip():
                return {"status": "error", "reason": "El archivo está vacío o no se pudo leer."}
            
            jd = user_job_desc if user_job_desc else "Perfil técnico general"
            decision = self.analyzer.analyze(raw_text, jd)
            
            # --- NUEVA LÓGICA DE EXTRACCIÓN (Score + Razón) ---
            f_score = 0
            reason = "No se pudo extraer una explicación detallada."
            texto_ia = str(decision)

            # 1. Extraer el score buscando específicamente la palabra "score" o el formato "XX%"
            match_score = re.search(r'score["\']?\s*[:\-]\s*(\d+)', texto_ia.lower())
            if match_score:
                f_score = int(match_score.group(1))
            else:
                # Si falla, busca números pero evita los de los formatos JSON
                numeros = re.findall(r'\b\d{2,3}\b', texto_ia) # Busca números de 2 o 3 cifras
                f_score = int(numeros[0]) if numeros else 0

            # 2.--- EXTRACCIÓN DEL MOTIVO (Jerarquía de Seguridad) ---
            # Plan A: Buscar formato JSON
            match_json = re.search(r'["\']motivo["\']\s*:\s*["\'](.*?)(?=["\']\s*[,\}])', texto_ia, re.DOTALL | re.IGNORECASE)
            # Plan B: Buscar texto natural (después de "motivo:")
            match_natural = re.search(r'(?:motivo|razón|explicación)\s*[:\-]\s*(.*)', texto_ia, re.DOTALL | re.IGNORECASE)

            if match_json:
                reason = match_json.group(1).strip()
            elif match_natural:
                reason = match_natural.group(1).strip()
            else:
                reason = texto_ia.strip()

            # Esto evita que se vean símbolos extraños en el Log de la pantalla
            reason = reason.replace('\\xa0', ' ').replace('\\n', '\n').replace('\\"', '"').replace('\\', '').strip()

            # --- Lógica de carpetas y Renombrado por Score ---
            nombre_original = os.path.basename(file_path)
            
            # Formateamos el score a 2 dígitos (ej: 05, 42, 98) para que el orden alfabético sea correcto
            score_prefix = f"{int(f_score):02d}"
            nuevo_nombre = f"{score_prefix}_{nombre_original}"

            # Verificamos si la IA puso "SI" en el campo apto del texto
            es_apto_por_texto = '"apto": "SI"' in texto_ia.upper() or "'apto': 'SI'" in texto_ia.upper()

            # Definición del destino según tu nueva escala (70, 50)
            if f_score >= 70 or es_apto_por_texto:
                destino = "RECLUTADOS"
            elif 50 <= f_score < 70:
                destino = "DUDAS"
            else:
                destino = "DESCARTADOS"

            # Ruta final con el NUEVO NOMBRE (incluye el score delante)
            ruta_final = os.path.join(self.base_output, destino, nuevo_nombre)
            
            if os.path.exists(file_path):
                # Usamos move pero con la ruta que contiene el nuevo_nombre
                shutil.move(file_path, ruta_final)

            self._append_to_report({
                "dest_path": ruta_final,
                "score": f_score,
                "reason": reason,
                "decision": destino
            })

            return {
                "status": "success",
                "decision": destino,
                "score": f_score,
                "reason": reason,  
                "dest_path": ruta_final
            }

        except Exception as e:
            logger.error(f"❌ Error procesando {file_path}: {str(e)}")
            return {"status": "error", "reason": str(e)}